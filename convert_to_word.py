import docx
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from docx.shared import Cm, Pt
from pathlib import Path


def change_page_margin(doc: docx.Document, margin: float):
    for section in doc.sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    

def change_line_spacing(doc: docx.Document):
    for i, paragraph in enumerate(doc.paragraphs):
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if (i - 5) % 7 == 0:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

def set_font_style(document: docx.Document):
    if 'Consola' not in document.styles:
        style = document.styles.add_style('Consola', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Consolas'
        font.size = Pt(10)
    for i, paragraph in enumerate(document.paragraphs):
        if (i - 5) % 7 == 0:
            paragraph.style = document.styles['Consola']
         
         
def set_cols(doc, n_cols):
    doc.add_section(WD_SECTION.CONTINUOUS) 
    section = doc.sections[-1]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(n_cols))
        

def create_intro(doc):
    doc.add_heading('TC de Programación', 0)
    doc.add_paragraph('Nombre: _____________________________________________________________________________________   No.: _____')
    p = doc.add_paragraph('')
    p.add_run('Para cada uno de los siguientes códigos en Python, diga su salida:').bold = True
    

def write_document(pathname):
    doc = docx.Document()
    for i, child in enumerate(Path(pathname).iterdir()):
        content = open(child, 'r').read()
        
        set_cols(doc, 1)
        create_intro(doc)    
        set_cols(doc, 2)
        
        doc.add_paragraph(content)
        doc.add_page_break()
        
    change_page_margin(doc, 1.20)
    change_line_spacing(doc)
    set_font_style(doc)
    doc.save('tests/Pruebas.docx')


if __name__ == '__main__':
    write_document('tests')